from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_title(doc: Document, text: str):
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return title


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc: Document, rows: int, cols: int, col_widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light List'
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def build_srs() -> Document:
    doc = Document()

    # Title Page
    add_title(doc, 'Software Requirements Specification (SRS)')
    add_para(doc, 'Project: AI-Powered Voice Shell (Multilingual)')
    add_para(doc, 'Version: 1.0')
    add_para(doc, 'Date: Auto-generated')
    add_para(doc, 'Authors: Project Team')

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    add_heading(doc, '1.1 Purpose', 2)
    add_para(doc, (
        'This SRS specifies the requirements for the AI-Powered Voice Shell. '
        'The system enables users to speak natural language (Tamil/English/etc.), '
        'which is transcribed and translated into a single, safe Linux command to be executed in WSL on Windows.'
    ))

    add_heading(doc, '1.2 Scope', 2)
    add_para(doc, (
        'The product captures voice, performs speech-to-text (Sarvam AI or Whisper), normalizes instructions using '
        'Google Gemini, enforces safety guardrails, executes the command in WSL, and provides audible and visual feedback.'
    ))

    add_heading(doc, '1.3 Definitions, Acronyms, and Abbreviations', 2)
    add_bullet(doc, 'STT: Speech-to-Text')
    add_bullet(doc, 'WSL: Windows Subsystem for Linux')
    add_bullet(doc, 'VAD: Voice Activity Detection')

    add_heading(doc, '1.4 References', 2)
    add_bullet(doc, 'Project README.md')
    add_bullet(doc, 'voice_shell_tunglish.py (main)')
    add_bullet(doc, 'terminal_gui.py (Rich TUI)')

    add_heading(doc, '1.5 Overview', 2)
    add_para(doc, 'Section 2 describes the overall product. Section 3 details specific requirements.')

    # 2. Overall Description
    add_heading(doc, '2. Overall Description', 1)
    add_heading(doc, '2.1 Product Perspective', 2)
    add_para(doc, (
        'Standalone desktop Python application running on Windows with WSL availability. '
        'Integrates external services: Sarvam AI (STT) and Google Gemini (LLM). Optional Whisper STT.'
    ))

    add_heading(doc, '2.2 Product Functions', 2)
    add_bullet(doc, 'Capture microphone audio with basic VAD and save as input.wav')
    add_bullet(doc, 'Transcribe speech using Sarvam AI (default) or Whisper')
    add_bullet(doc, 'Normalize/translate to clear English and generate one Linux command using Gemini')
    add_bullet(doc, 'Sanitize command to block interactive tools and risky patterns')
    add_bullet(doc, 'Execute in WSL while persisting working directory across cd')
    add_bullet(doc, 'Provide console logs and spoken feedback (pyttsx3)')

    add_heading(doc, '2.3 User Classes and Characteristics', 2)
    add_bullet(doc, 'Developer/Power user on Windows with WSL installed')
    add_bullet(doc, 'Users comfortable with basic shell operations; prefers voice interaction')

    add_heading(doc, '2.4 Operating Environment', 2)
    add_bullet(doc, 'Windows 10/11, Python 3.10+')
    add_bullet(doc, 'WSL installed and accessible via "wsl" on PATH')
    add_bullet(doc, 'Microphone access and audio devices configured')

    add_heading(doc, '2.5 Design and Implementation Constraints', 2)
    add_bullet(doc, 'Non-interactive commands only; editors/pagers are blocked')
    add_bullet(doc, 'Per-command timeout (default 10s)')
    add_bullet(doc, 'Network connectivity required for Sarvam AI and Gemini')

    add_heading(doc, '2.6 Assumptions and Dependencies', 2)
    add_bullet(doc, 'Valid SARVAM_API_KEY configured in .env')
    add_bullet(doc, 'Gemini API key configured (currently in code)')
    add_bullet(doc, 'Python dependencies per requirements.txt installed')

    # 3. External Interface Requirements
    add_heading(doc, '3. External Interface Requirements', 1)
    add_heading(doc, '3.1 User Interfaces', 2)
    add_bullet(doc, 'Console CLI prompts and outputs')
    add_bullet(doc, 'Optional Rich TUI (terminal_gui.py) with live panels for transcript, command, and output')

    add_heading(doc, '3.2 Hardware Interfaces', 2)
    add_bullet(doc, 'Microphone input device')

    add_heading(doc, '3.3 Software Interfaces', 2)
    add_bullet(doc, 'Sarvam AI SDK (sarvamai) for STT')
    add_bullet(doc, 'Google Generative AI (Gemini) for normalization and command generation')
    add_bullet(doc, 'WSL subsystem for Linux command execution')

    add_heading(doc, '3.4 Communications Interfaces', 2)
    add_bullet(doc, 'HTTPS to Sarvam AI and Google Gemini APIs')

    # 4. System Features (Functional Requirements)
    add_heading(doc, '4. System Features', 1)

    add_heading(doc, '4.1 Voice Capture and VAD', 2)
    add_numbered(doc, 'The system shall record audio from the default microphone until silence is detected.')
    add_numbered(doc, 'The system shall save recordings as input.wav with preroll and tail frames.')

    add_heading(doc, '4.2 Speech Transcription', 2)
    add_numbered(doc, 'The system shall transcribe speech using Sarvam AI by default.')
    add_numbered(doc, 'The system shall support Whisper as an alternative provider.')
    add_numbered(doc, 'The system shall allow language and model configuration via .env.')

    add_heading(doc, '4.3 Command Generation and Normalization', 2)
    add_numbered(doc, 'The system shall normalize instructions into clear English using Gemini.')
    add_numbered(doc, 'The system shall generate exactly one non-interactive Linux command in JSON output.')
    add_numbered(doc, 'The system shall detect exit intent and signal termination when confirmed.')

    add_heading(doc, '4.4 Safety Guardrails', 2)
    add_numbered(doc, 'The system shall block interactive tools (nano, vi, vim, emacs, ed, less, more, man, top, htop, watch, code, notepad).')
    add_numbered(doc, 'The system shall block interactive redirection patterns (e.g., cat > file).')

    add_heading(doc, '4.5 Command Execution in WSL', 2)
    add_numbered(doc, 'The system shall execute commands via wsl --cd <cwd> -- bash -lc <command>.')
    add_numbered(doc, 'The system shall persist the working directory across cd operations.')
    add_numbered(doc, 'The system shall enforce a per-command timeout (default 10 seconds).')

    add_heading(doc, '4.6 Feedback and Logging', 2)
    add_numbered(doc, 'The system shall print command, stdout, and stderr to the console or TUI.')
    add_numbered(doc, 'The system shall provide spoken feedback using text-to-speech.')

    # 5. Nonfunctional Requirements
    add_heading(doc, '5. Other Nonfunctional Requirements', 1)

    add_heading(doc, '5.1 Performance Requirements', 2)
    add_bullet(doc, 'Transcription and command generation should complete within a few seconds for typical inputs.')

    add_heading(doc, '5.2 Security Requirements', 2)
    add_bullet(doc, 'API keys shall not be logged. Keys shall be stored in environment (.env) or secure key vaults.')

    add_heading(doc, '5.3 Reliability and Availability', 2)
    add_bullet(doc, 'Graceful handling of network failures and API errors with clear messages.')

    add_heading(doc, '5.4 Usability', 2)
    add_bullet(doc, 'Simple prompts, clear feedback, and minimal configuration steps.')

    add_heading(doc, '5.5 Portability', 2)
    add_bullet(doc, 'Runs on Windows 10/11 with WSL and Python 3.10+. Optional Linux adjustments may be needed if ported.')

    # 6. Other Requirements
    add_heading(doc, '6. Other Requirements', 1)
    add_bullet(doc, 'Compliance: Follow provider usage policies for Sarvam AI and Google Gemini.')
    add_bullet(doc, 'Logging: Avoid storing sensitive audio/text by default; debug logs should be optional.')

    # 7. Use Cases
    add_heading(doc, '7. Use Cases', 1)
    add_heading(doc, 'UC-1: List files', 2)
    add_bullet(doc, 'Actor: User')
    add_bullet(doc, 'Preconditions: App running; microphone available; WSL installed')
    add_bullet(doc, 'Main Flow: Speak "List files here" → Transcribe → Generate "ls -la" → Execute → Output listed')

    add_heading(doc, 'UC-2: Create folder and navigate', 2)
    add_bullet(doc, 'Main Flow: "Sample folder create pannu" → mkdir Sample; "Change directory to Sample" → cd Sample')

    add_heading(doc, 'UC-3: View file', 2)
    add_bullet(doc, 'Main Flow: "Open sample.py file" → cat sample.py')

    # 8. Acceptance Criteria (Examples)
    add_heading(doc, '8. Acceptance Criteria (Samples)', 1)
    add_bullet(doc, 'Given the app is running, when the user says "Exit", then the app terminates gracefully.')
    add_bullet(doc, 'Given poor network, when Sarvam/Gemini fails, then the app shows an error and does not crash.')
    add_bullet(doc, 'Given a request to use nano, when detected, then the app blocks the command and informs the user.')

    # 9. Future Enhancements
    add_heading(doc, '9. Future Enhancements', 1)
    add_bullet(doc, 'Command confirmation for sensitive actions')
    add_bullet(doc, 'Improved Tamil prompt tuning and custom vocabulary')
    add_bullet(doc, 'History, undo helpers, and safer rm patterns')
    add_bullet(doc, 'Multi-step task orchestration with approvals')

    # 10. Appendix
    add_heading(doc, '10. Appendix', 1)
    add_para(doc, 'Environment Variables:')
    add_bullet(doc, 'SARVAM_API_KEY, SARVAM_STT_MODEL, SARVAM_LANGUAGE_CODE')
    add_bullet(doc, 'TRANSCRIBE_PROVIDER (sarvam|whisper)')
    add_bullet(doc, 'VOICE_SHELL_WSL_DISTRO, VOICE_SHELL_WSL_USER (optional)')

    return doc


if __name__ == '__main__':
    document = build_srs()
    output_name = 'AI_Powered_Voice_Shell_SRS.docx'
    document.save(output_name)
    print(f'Saved {output_name}')
