from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)
    return p


def build_report() -> Document:
    doc = Document()

    # Title
    title = doc.add_heading("AI-Powered Voice Shell (Multilingual) — Project Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata (optional)
    add_para(doc, "Platform: Windows + WSL | Core: Sarvam AI STT, Google Gemini, Python")

    # Abstract
    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        (
            "This project enables hands-free control of a Linux shell on Windows. A user speaks a command "
            "(Tamil/English/etc.), which is recorded, transcribed using Sarvam AI Speech-to-Text, translated/"
            "normalized by Gemini to clear English, and converted into a safe, non-interactive Linux command "
            "executed within WSL. The system reads out status and prints outputs, providing an accessible and "
            "productivity-oriented interface for common CLI tasks."
        ),
    )

    # Introduction
    add_heading(doc, "Introduction", 1)
    add_para(
        doc,
        (
            "Traditional shell usage demands keyboard proficiency and context switching. This assistant bridges "
            "natural language and shell commands, offering voice-first interaction, multilingual input targeting "
            "Tamil and English, safety guardrails (no interactive tools), and seamless execution inside WSL for "
            "Windows developers."
        ),
    )
    add_para(
        doc,
        (
            "Use cases include quick file operations, directory navigation, inspecting logs, and simple DevOps "
            "queries without leaving your current focus."
        ),
    )

    # Implementation
    add_heading(doc, "Implementation", 1)
    add_para(
        doc,
        (
            "Audio capture and VAD: record_voice() uses sounddevice with basic voice-activity detection, saving "
            "input.wav with pre-roll and tail."
        ),
    )
    add_para(
        doc,
        (
            "Transcription (STT): Default is Sarvam AI SDK (sarvamai). The app calls SarvamAI.speech_to_text.transcribe "
            "with model and language_code from .env. Whisper (faster-whisper) remains optional."
        ),
    )
    add_para(
        doc,
        (
            "Language normalization and command generation: Gemini rewrites instructions to clear English and returns "
            "a strict-JSON response with a single non-interactive Linux command and an exit flag."
        ),
    )
    add_para(
        doc,
        (
            "Safety and sanitization: Interactive tools (nano, vim, less, man, top, watch, code, etc.) and risky "
            "patterns (e.g., 'cat > file') are blocked; commands have timeouts."
        ),
    )
    add_para(
        doc,
        (
            "WSL execution: Commands run via 'wsl --cd <cwd> -- bash -lc <command>'. The initial WSL path maps from the "
            "Windows project folder (/mnt/c/...). The app tracks 'cd' to persist working directory. Optional env vars "
            "VOICE_SHELL_WSL_DISTRO and VOICE_SHELL_WSL_USER select the distro/user."
        ),
    )
    add_para(
        doc,
        (
            "Feedback: Console prints show command, outputs, errors, and WSL user/distro; pyttsx3 provides spoken status."
        ),
    )
    add_para(
        doc,
        (
            "Configuration: .env (python-dotenv) holds SARVAM_API_KEY, SARVAM_STT_MODEL, SARVAM_LANGUAGE_CODE, and "
            "TRANSCRIBE_PROVIDER; Gemini key is configured in the code."
        ),
    )

    # Working
    add_heading(doc, "Working", 1)
    add_para(
        doc,
        (
            "Flow: Speak → record (VAD) → input.wav → Sarvam STT → text → Gemini normalization + command generation → "
            "sanitize → execute in WSL → outputs + TTS → loop or exit."
        ),
    )
    add_para(
        doc,
        (
            "Examples: 'Sample folder create pannu' → mkdir Sample; 'Open sample.py file' → cat sample.py; "
            "'Change to Sample' → cd Sample; 'Exit' → terminate."
        ),
    )

    # Limitations
    add_heading(doc, "Limitations", 1)
    add_para(
        doc,
        (
            "STT accuracy depends on language_code, model, accent, and audio quality. Gemini may sometimes propose "
            "context-lacking commands; interactive programs are intentionally blocked; long-running tasks may hit "
            "timeouts; and WSL is required for execution."
        ),
    )

    # Conclusion
    add_heading(doc, "Conclusion", 1)
    add_para(
        doc,
        (
            "This project demonstrates a practical, multilingual, voice-driven interface to the Linux shell on Windows. "
            "By combining Sarvam AI for speech recognition, Gemini for command generation, and WSL for execution, it "
            "enables accessible and hands-free system interactions with safety guardrails. With targeted tuning and "
            "prompt improvements, it can serve as a robust productivity tool in everyday development."
        ),
    )

    return doc


if __name__ == "__main__":
    document = build_report()
    output_name = "Project_Report.docx"
    document.save(output_name)
    print(f"Saved {output_name}")
